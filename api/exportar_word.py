from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from helpers import get_feedback, get_form_by_homework
import json
import time
import io
import re

# Función para crear hipervínculos en un párrafo
def add_hyperlink(paragraph, text, url, relationship_id):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)  # Establecer el rId para el hipervínculo

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")  # Propiedades del "run"
    # Color y subrayado para el hipervínculo
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")  # Azul
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")  # Subrayado
    rPr.append(underline)

    run.append(rPr)

    text_element = OxmlElement("w:t")
    text_element.text = text  # Texto para el hipervínculo
    run.append(text_element)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

# Función para imprimir el documento con hipervínculos
def document_print(dic_list):
    doc = Document()

    # Para cada diccionario en la lista
    for item in dic_list:
        for key, value in item.items():
            text = str(value[1])

            # Buscar URLs en el texto
            url_pattern = r"http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+"
            urls = re.findall(url_pattern, text)

            # Crear el párrafo
            if value[0] == 2:
                doc.add_heading(text, level=1)
            else:
                if value[0] == 0:
                    text = text
                    paragraph = doc.add_paragraph()
                elif value[0] == 1:
                    text = '\n' + text
                    paragraph = doc.add_paragraph()
                key_run = paragraph.add_run(f"{key}: ")
                key_run.bold = True

                # Dividir el texto por las URLs para insertar hipervínculos
                partes = re.split(url_pattern, text)

                # Agregar el texto y los hipervínculos
                for i, parte in enumerate(partes):
                    paragraph.add_run(parte)  # Añadir texto normal
                    if i < len(urls):
                        # Crear relación para el hipervínculo
                        part = doc._part
                        relationship_id = part.relate_to(urls[i], RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                        add_hyperlink(paragraph, urls[i], urls[i], relationship_id)

        doc.add_paragraph()  # Añadir espacio entre secciones

    # Guardar el documento en un objeto BytesIO
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)

    return f  # Devolver el objeto BytesIO

# Función para preparar el diccionario
def preparar_diccionario(request_list, link_form_homework):
    dic_lis = []
    for request_key in request_list.keys():
        request = request_list[request_key]
        dic = {}
        dic['Archivo Nombre'] = (2, request['file_name'])
        dic['Fecha Proceso'] = (0, request['time_stamp'])
        dic['Tarea Entregada'] = (1, request['user_prompt'])

        feedback = get_feedback(request['result_text'])
        feedback += '\n' + '\nHelp us to improve! Rate the feedback given by filling the following survey:'
        if ('homework_number' in request.keys()):
            feedback += '\n' + get_form_by_homework(request['homework_number'])
        else:
            feedback += '\n' + get_form_by_homework('x')
        dic['Feedback'] = (1, feedback)
        dic_lis.append(dic)

    # Generar el documento de Word con la función document_print
    doc_file = document_print(dic_lis)

    return doc_file  # Devolver el objeto BytesIO
